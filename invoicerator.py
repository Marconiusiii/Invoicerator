from __future__ import annotations

import argparse
import importlib.util
import json
from dataclasses import dataclass
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DEFAULT_FONT = "Helvetica"
DEFAULT_DUE_DAYS = 30


@dataclass(frozen=True)
class UserProfile:
    name: str
    street_address: str
    city_state_zip: str
    phone: str
    email: str
    late_fees: str = ""
    payment_instructions: str = ""


@dataclass(frozen=True)
class InvoiceEntry:
    service_date: str
    project: str
    hours: Decimal


@dataclass(frozen=True)
class InvoiceData:
    client: str
    invoice_number: str
    submitted_date: str
    due_date: str
    hourly_rate: Decimal
    entries: list[InvoiceEntry]

    @property
    def total_hours(self) -> Decimal:
        return sum((entry.hours for entry in self.entries), Decimal("0"))

    @property
    def total_owed(self) -> Decimal:
        return self.total_hours * self.hourly_rate


class InvoiceratorError(RuntimeError):
    pass


def prompt_text(label: str, *, required: bool = True, default: Optional[str] = None) -> str:
    while True:
        raw = input(f"{label}: ").strip()
        if raw:
            return raw
        if default is not None:
            return default
        if not required:
            return ""
        print("Please enter a value.")


def prompt_date(label: str, *, default: Optional[str] = None) -> str:
    while True:
        value = prompt_text(label, default=default)
        try:
            datetime.strptime(value, "%m/%d/%Y")
        except ValueError:
            print("Use MM/DD/YYYY format (example: 03/05/2026).")
            continue
        return value


def prompt_decimal(label: str) -> Decimal:
    while True:
        value = input(f"{label}: ").strip()
        try:
            number = Decimal(value)
        except InvalidOperation:
            print("Enter a valid number (example: 75 or 75.00).")
            continue

        if number < 0:
            print("Value cannot be negative.")
            continue

        return number


def prompt_quarter_hours(label: str) -> Decimal:
    while True:
        hours = prompt_decimal(label)
        if (hours * 100) % 25 != 0:
            print("Hours must be in 0.25 increments (example: 1.25).")
            continue
        return hours


def yes_no(label: str, *, default: bool = False) -> bool:
    suffix = "[Y/n]" if default else "[y/N]"
    while True:
        value = input(f"{label} {suffix}: ").strip().lower()
        if not value:
            return default
        if value in {"y", "yes"}:
            return True
        if value in {"n", "no"}:
            return False
        print("Enter y or n.")


def format_hours(value: Decimal) -> str:
    return f"{value:.2f}".rstrip("0").rstrip(".")


def apply_run_style(run, *, bold: bool = False, size: int = 12) -> None:
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(size)
    run.bold = bold


def load_profile_from_json(path: Path) -> UserProfile:
    with path.open("r", encoding="utf-8") as handle:
        payload = json.load(handle)

    def field(name: str) -> str:
        value = str(payload.get(name, "")).strip()
        if not value and name in {"name", "street_address", "city_state_zip", "phone", "email"}:
            raise InvoiceratorError(f"Missing required profile field: {name}")
        return value

    return UserProfile(
        name=field("name"),
        street_address=field("street_address"),
        city_state_zip=field("city_state_zip"),
        phone=field("phone"),
        email=field("email"),
        late_fees=field("late_fees"),
        payment_instructions=field("payment_instructions"),
    )


def load_profile_from_legacy_module(path: Path) -> UserProfile:
    spec = importlib.util.spec_from_file_location("legacy_user_data", str(path))
    if spec is None or spec.loader is None:
        raise InvoiceratorError(f"Could not load legacy profile at {path}")

    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    def attr(name: str, *, required: bool = False) -> str:
        value = str(getattr(module, name, "")).strip()
        if required and not value:
            raise InvoiceratorError(f"Missing required legacy profile value: {name}")
        return value

    return UserProfile(
        name=attr("userName", required=True),
        street_address=attr("userStreetAddress", required=True),
        city_state_zip=attr("userCityStateZip", required=True),
        phone=attr("userPhone", required=True),
        email=attr("userEmail", required=True),
        late_fees=attr("userLateFees"),
        payment_instructions=attr("userPay"),
    )


def load_profile(profile_path: Optional[Path]) -> UserProfile:
    if profile_path:
        if profile_path.suffix.lower() == ".json":
            return load_profile_from_json(profile_path)
        if profile_path.suffix.lower() == ".py":
            return load_profile_from_legacy_module(profile_path)
        raise InvoiceratorError("Profile file must be .json or .py")

    default_json = Path("user_profile.json")
    if default_json.exists():
        return load_profile_from_json(default_json)

    legacy_py = Path("userData.py")
    if legacy_py.exists():
        return load_profile_from_legacy_module(legacy_py)

    raise InvoiceratorError(
        "No profile found. Add user_profile.json (recommended) or userData.py (legacy)."
    )


def gather_invoice_data() -> InvoiceData:
    today = datetime.now().strftime("%m/%d/%Y")

    client = prompt_text("Client name")
    invoice_number = prompt_text("Invoice number")

    entries: list[InvoiceEntry] = []
    while True:
        print("\nAdd invoice line item")
        service_date = prompt_date("Date of service (MM/DD/YYYY)", default=today)
        project = prompt_text("Project description")
        hours = prompt_quarter_hours("Hours worked (0.25 increments)")

        entry = InvoiceEntry(service_date=service_date, project=project, hours=hours)
        print(f"Entry: {entry.service_date} | {entry.project} | {format_hours(entry.hours)}")

        if yes_no("Keep this entry?", default=True):
            entries.append(entry)

        if not yes_no("Add another line item?", default=False):
            break

    if not entries:
        raise InvoiceratorError("At least one line item is required.")

    hourly_rate = prompt_decimal("Hourly rate")
    submitted_date = prompt_date("Date submitted", default=today)
    submitted_dt = datetime.strptime(submitted_date, "%m/%d/%Y")
    due_date = (submitted_dt + timedelta(days=DEFAULT_DUE_DAYS)).strftime("%m/%d/%Y")
    print(f"Payment due date ({DEFAULT_DUE_DAYS} days): {due_date}")

    return InvoiceData(
        client=client,
        invoice_number=invoice_number,
        submitted_date=submitted_date,
        due_date=due_date,
        hourly_rate=hourly_rate,
        entries=entries,
    )


def write_invoice_docx(profile: UserProfile, invoice: InvoiceData, out_path: Path) -> None:
    document = Document()

    title = document.add_heading().add_run(profile.name)
    apply_run_style(title, size=18)

    header = document.add_paragraph()
    header_text = (
        f"{profile.street_address}\n"
        f"{profile.city_state_zip}\n"
        f"{profile.phone}\n"
        f"{profile.email}\n"
    )
    header_run = header.add_run(header_text)
    apply_run_style(header_run)

    document.add_paragraph()

    invoice_heading = document.add_heading("INVOICE", level=2)
    invoice_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    metadata = document.add_paragraph()
    in_to = metadata.add_run("Invoicing to: ")
    apply_run_style(in_to, bold=True)
    in_to_val = metadata.add_run(invoice.client)
    apply_run_style(in_to_val)

    metadata.add_run("\n")

    inv_num = metadata.add_run("Invoice #: ")
    apply_run_style(inv_num, bold=True)
    inv_num_val = metadata.add_run(invoice.invoice_number)
    apply_run_style(inv_num_val)

    table = document.add_table(rows=1, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Date"
    header_cells[1].text = "Project"
    header_cells[2].text = "Hours"

    for entry in invoice.entries:
        row = table.add_row().cells
        row[0].text = entry.service_date
        row[1].text = entry.project
        row[2].text = format_hours(entry.hours)

    rate_p = document.add_paragraph()
    rt = rate_p.add_run("Rate: ")
    apply_run_style(rt, bold=True)
    rt_val = rate_p.add_run(f"${invoice.hourly_rate:,.2f}")
    apply_run_style(rt_val)

    hours_p = document.add_paragraph()
    total_hours = hours_p.add_run("Total Hours: ")
    apply_run_style(total_hours, bold=True)
    total_hours_val = hours_p.add_run(format_hours(invoice.total_hours))
    apply_run_style(total_hours_val)

    total_p = document.add_paragraph()
    total = total_p.add_run("Total Owed: ")
    apply_run_style(total, bold=True)
    total_val = total_p.add_run(f"${invoice.total_owed:,.2f}")
    apply_run_style(total_val)

    date_sub = document.add_paragraph()
    submitted = date_sub.add_run(f"Date Submitted: {invoice.submitted_date}")
    apply_run_style(submitted)

    pay_date = document.add_paragraph()
    pay_due = pay_date.add_run(f"Payment Due: {invoice.due_date}")
    apply_run_style(pay_due)

    if profile.late_fees:
        late = document.add_paragraph().add_run(profile.late_fees)
        apply_run_style(late)

    if profile.payment_instructions:
        how_to_pay = document.add_paragraph().add_run(profile.payment_instructions)
        apply_run_style(how_to_pay)

    document.save(out_path)


def parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate DOCX invoices from Terminal")
    parser.add_argument(
        "--profile",
        type=Path,
        default=None,
        help="Path to profile file (.json recommended, .py legacy supported)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Output DOCX path. If omitted, you'll be prompted.",
    )
    return parser.parse_args(argv)


def ensure_docx_suffix(path: Path) -> Path:
    if path.suffix.lower() == ".docx":
        return path
    return path.with_suffix(".docx")


def main(argv: Optional[Iterable[str]] = None) -> int:
    try:
        args = parse_args(argv)

        print("Invoicerator 3.0")
        print("Generate invoice DOCX files from Terminal.")

        profile = load_profile(args.profile)
        invoice = gather_invoice_data()

        out_path = args.output
        if out_path is None:
            out_path = Path(prompt_text("Output filename (.docx appended automatically if needed)"))

        out_path = ensure_docx_suffix(out_path)
        write_invoice_docx(profile, invoice, out_path)

        print(f"Saved DOCX: {out_path}")
        print(f"Total owed: ${invoice.total_owed:,.2f}")
        return 0
    except InvoiceratorError as exc:
        print(f"Error: {exc}")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
